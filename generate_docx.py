import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()
doc.add_heading('Tabel Adopsi Lisensi - NutriVision AI', 0)

p = doc.add_paragraph('Berikut adalah rincian pustaka (library), kerangka kerja (framework), dan perangkat lunak pihak ketiga yang digunakan dalam pengembangan aplikasi NutriVision AI, disusun sesuai format adopsi lisensi GEMASTIK.')

table_data = [
    ("No", "Nama Komponen / Library", "Sumber / Pembuat", "Jenis Lisensi", "Keterangan Penggunaan"),
    ("1", "React.js", "Meta Open Source (https://react.dev)", "MIT License", "Kerangka kerja (framework) utama untuk membangun antarmuka pengguna interaktif pada sisi Frontend."),
    ("2", "Tailwind CSS", "Tailwind Labs (https://tailwindcss.com)", "MIT License", "Kerangka kerja CSS untuk mempercepat proses desain tata letak, warna, dan responsivitas UI aplikasi."),
    ("3", "Vite", "Evan You & Komunitas Vite (https://vitejs.dev)", "MIT License", "Build tool dan peladen pengembangan (development server) untuk melakukan kompilasi Frontend dengan cepat."),
    ("4", "React Router", "Remix Software (https://reactrouter.com)", "MIT License", "Modul untuk menangani navigasi dan perpindahan rute antar halaman secara mandiri di sisi klien (Single Page Application)."),
    ("5", "Express.js", "OpenJS Foundation (https://expressjs.com)", "MIT License", "Kerangka kerja backend berbasis Node.js yang digunakan sebagai server utama pengelola Application Programming Interface (API)."),
    ("6", "Prisma ORM", "Prisma Data, Inc. (https://www.prisma.io)", "Apache 2.0", "Object-Relational Mapping (ORM) yang mempermudah interaksi dan pengelolaan skema data ke dalam database."),
    ("7", "pg (node-postgres)", "Brian Carlson & Komunitas (https://node-postgres.com)", "MIT License", "Driver koneksi yang menjembatani server backend Node.js dengan sistem database PostgreSQL."),
    ("8", "FastAPI", "Sebastián Ramírez (https://fastapi.tiangolo.com)", "MIT License", "Kerangka kerja Python berkinerja tinggi yang menjalankan modul AI Engine untuk pemrosesan analitik citra (foto)."),
    ("9", "Uvicorn", "Encode OSS (https://www.uvicorn.org)", "BSD-3-Clause", "Peladen web (web server) ASGI yang bertugas mengeksekusi secara asinkron aplikasi AI Engine (FastAPI)."),
    ("10", "Python-multipart", "Andrew L.F. (https://github.com/andrew-d/python-multipart)", "Apache 2.0", "Pustaka pendukung pada AI Engine yang difungsikan untuk menerima unggahan file gambar dari frontend melalui form-data."),
    ("11", "Dotenv", "Scott Motte (https://github.com/motdotla/dotenv)", "BSD-2-Clause", "Modul ringan yang bertugas untuk memuat variabel lingkungan dan rahasia konfigurasi (seperti password database) ke dalam sistem.")
]

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
for i, heading in enumerate(table_data[0]):
    hdr_cells[i].text = heading
    for paragraph in hdr_cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

# Data rows
for row_data in table_data[1:]:
    row_cells = table.add_row().cells
    for i, cell_text in enumerate(row_data):
        row_cells[i].text = cell_text

note = doc.add_paragraph('\nCatatan Legalitas: Seluruh pustaka di atas menggunakan lisensi open-source permisif (MIT, Apache 2.0, dan BSD) yang mengizinkan pemanfaatan, modifikasi, serta komersialisasi perangkat lunak tanpa mewajibkan tim pengembang untuk mempublikasikan kode sumber (lisensi bebas/terbuka yang sah).')

save_path = os.path.join(os.getcwd(), 'Adopsi_Lisensi_NutriVision_AI.docx')
doc.save(save_path)
print(f"Saved doc to {save_path}")
